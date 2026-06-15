from __future__ import annotations

from collections.abc import Iterable
from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    ListFlowable,
    ListItem,
    PageTemplate,
    Paragraph,
    Spacer,
)


@dataclass(frozen=True)
class ExportSection:
    name: str
    content_json: dict[str, Any]


def _text_content(node: dict[str, Any]) -> str:
    if node.get("type") == "text":
        return str(node.get("text", ""))
    return "".join(
        _text_content(child)
        for child in node.get("content", [])
        if isinstance(child, dict)
    )


def _iter_blocks(document: dict[str, Any]) -> Iterable[tuple[str, str, int | None]]:
    for node in document.get("content", []):
        if not isinstance(node, dict):
            continue
        node_type = node.get("type")
        if node_type == "heading":
            level = node.get("attrs", {}).get("level", 2)
            yield ("heading", _text_content(node), level if isinstance(level, int) else 2)
        elif node_type in {"paragraph", "blockquote"}:
            yield (node_type, _text_content(node), None)
        elif node_type in {"bulletList", "orderedList"}:
            for item in node.get("content", []):
                if isinstance(item, dict):
                    yield (node_type, _text_content(item), None)


def _pdf_header_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFillColor(colors.HexColor("#475569"))
    canvas.setFont("Helvetica", 8)
    canvas.drawString(20 * mm, 12 * mm, "EuroGrant AI")
    canvas.drawRightString(A4[0] - 20 * mm, 12 * mm, f"Page {doc.page}")
    canvas.restoreState()


def generate_pdf(
    proposal_id: int,
    grant_title: str,
    sections: list[ExportSection],
) -> bytes:
    buffer = BytesIO()
    document = BaseDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=22 * mm,
        bottomMargin=22 * mm,
        title=f"Proposal {proposal_id} - {grant_title}",
        author="EuroGrant AI",
    )
    frame = Frame(document.leftMargin, document.bottomMargin, document.width, document.height)
    document.addPageTemplates(
        [PageTemplate(id="proposal", frames=[frame], onPage=_pdf_header_footer)]
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ProposalTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=28,
        textColor=colors.HexColor("#064E3B"),
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    section_style = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=21,
        textColor=colors.HexColor("#065F46"),
        spaceBefore=12,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "ProposalBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=16,
        textColor=colors.HexColor("#1F2937"),
        spaceAfter=8,
    )
    quote_style = ParagraphStyle(
        "ProposalQuote",
        parent=body_style,
        leftIndent=10,
        borderColor=colors.HexColor("#B45309"),
        borderWidth=1,
        borderPadding=6,
    )

    story: list[Any] = [
        Paragraph("Grant Proposal", title_style),
        Paragraph(f"Target grant: {grant_title}", styles["Heading2"]),
        Paragraph(f"Proposal reference: {proposal_id}", body_style),
        Spacer(1, 12),
    ]
    for section in sections:
        story.append(Paragraph(_escape_xml(section.name), section_style))
        for block_type, text, level in _iter_blocks(section.content_json):
            if not text.strip():
                story.append(Spacer(1, 6))
                continue
            safe_text = _escape_xml(text)
            if block_type == "heading":
                style = styles["Heading2"] if (level or 2) <= 2 else styles["Heading3"]
                story.append(Paragraph(safe_text, style))
            elif block_type == "blockquote":
                story.append(Paragraph(safe_text, quote_style))
            elif block_type in {"bulletList", "orderedList"}:
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(safe_text, body_style))],
                        bulletType="bullet" if block_type == "bulletList" else "1",
                        leftIndent=18,
                    )
                )
            else:
                story.append(Paragraph(safe_text, body_style))

    document.build(story)
    return buffer.getvalue()


def generate_docx(
    proposal_id: int,
    grant_title: str,
    sections: list[ExportSection],
) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_heading("Grant Proposal", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Target grant: {grant_title}")
    document.add_paragraph(f"Proposal reference: {proposal_id}")

    for export_section in sections:
        document.add_heading(export_section.name, level=1)
        for block_type, text, level in _iter_blocks(export_section.content_json):
            if block_type == "heading":
                document.add_heading(text, level=max(2, min(level or 2, 3)))
            elif block_type == "bulletList":
                document.add_paragraph(text, style="List Bullet")
            elif block_type == "orderedList":
                document.add_paragraph(text, style="List Number")
            elif block_type == "blockquote":
                paragraph = document.add_paragraph(text)
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.style = document.styles["Quote"]
            else:
                document.add_paragraph(text)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(31, 41, 55)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _escape_xml(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
